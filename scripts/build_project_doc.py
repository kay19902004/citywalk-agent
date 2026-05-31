from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path("/Users/a/Desktop/作品/citywalk-agent")
SCREENSHOT_OUT = ROOT / "docs" / "project-screenshots"
DOCX_OUT = ROOT / "docs" / "CityWalk_现实城市冒险项目说明.docx"


PAGES = [
    {
        "title": "生成页",
        "route": "/",
        "file": "01-home.png",
        "description": "选择城市、故事气质、天气和时间，生成一局包含案卷、身份卡与 5 站路线的现实城市冒险。",
    },
    {
        "title": "位置页",
        "route": "/location",
        "file": "02-location.png",
        "description": "补充 GPS、地标、街区类型和现场观察。",
    },
    {
        "title": "案卷页",
        "route": "/dossier",
        "file": "03-dossier.png",
        "description": "展示长开场、事件概述、未解问题、路线预告和 3 张可选身份卡。",
    },
    {
        "title": "剧情页",
        "route": "/play",
        "file": "04-play.png",
        "description": "承载当前站点目标、身份探索加成、5 站路线进度卡、NPC 信息和拍照入口。",
    },
    {
        "title": "拍照页",
        "route": "/photo",
        "file": "05-photo.png",
        "description": "产品高光页：未上传照片时展示 HUD 示例，上传照片后叠加主线标记、隐藏标记和方向箭头。",
    },
    {
        "title": "线索页",
        "route": "/clues",
        "file": "06-clues.png",
        "description": "汇总身份、人物、证物、时间线、照片发现；完成 5 站后显示冒险报告和分享卡。",
    },
]


def main() -> None:
    SCREENSHOT_OUT.mkdir(parents=True, exist_ok=True)
    normalized = normalize_screenshots()

    doc = Document()
    configure_document(doc)

    add_title(doc)
    add_overview(doc)
    add_logic(doc)
    add_architecture(doc)
    add_ai_flow(doc)
    add_page_inventory(doc)
    add_screenshots(doc, normalized)
    add_next_steps(doc)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


def normalize_screenshots() -> dict[str, Path]:
    normalized: dict[str, Path] = {}
    for page in PAGES:
        src = SCREENSHOT_OUT / page["file"]
        out = SCREENSHOT_OUT / page["file"]
        with Image.open(src) as image:
            image.convert("RGB").save(out, "PNG")
        normalized[page["file"]] = out
    return normalized


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "CityWalk 现实城市冒险 Agent"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.text = "项目说明与页面截图"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("CityWalk 现实城市冒险 Agent 项目说明")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.add_run(f"生成日期：{date.today().isoformat()}  |  当前版本：HUD 照片驱动城市冒险 MVP")

    add_callout(
        doc,
        "一句话概括",
        "这个项目把传统文本式 CityWalk 剧情改造成现实城市冒险：玩家先生成一条五站路线，到真实地点拍照，AI 根据照片生成主线任务、隐藏任务和可点击 HUD 线索，让现实街区变成轻量 AR 游戏界面。",
    )


def add_overview(doc: Document) -> None:
    doc.add_heading("1. 项目概览", level=1)
    doc.add_paragraph(
        "当前项目是一个基于 Next.js App Router 的单人现实城市冒险 Agent。它不再强调传统剧本杀的封闭案卷，而是把用户所在城市、天气、时间、地标和现场观察转化成可以步行完成的 5 站路线。"
    )
    doc.add_paragraph(
        "核心体验是“生成故事 -> 选择身份 -> 到站拍照 -> 点击照片标记 -> 解锁线索 -> 推进下一站”。用户可以把街道、店招、路牌、橱窗、地铁口、公共装置等真实元素变成任务锚点。当前 MVP 的重点是照片页的 HUD 爽点：用户不用理解复杂剧情，也能一眼看懂“拍下街景后，线索会浮现在照片上”。"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [2200, 3200, 3960])
    headers = ["模块", "作用", "当前状态"]
    for index, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], text, bold=True, fill="F2F4F7")

    rows = [
        ("故事生成", "DeepSeek 生成长开场、案卷、身份卡、5 站路线和节点任务。", "已接入，包含 JSON 截断重试和质量归一化。"),
        ("现实照片", "Qwen-VL 分析现场照片，输出 HUD 标记、隐藏任务和方向箭头。", "已接入，失败时显示本地临时线索。"),
        ("游戏推进", "主线标记推进站点，隐藏标记只加入线索簿。", "已实现主线/隐藏双轨推进。"),
        ("UI 流程", "底部导航覆盖生成、位置、案卷、剧情、拍照、线索。", "已完成 MVP 页面与截图。"),
        ("安全边界", "任务只允许公共可见观察，不进入私人区域、不拍陌生人、不接触物品。", "已加入 normalize / quality guard。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            set_cell_text(cells[index], text)


def add_logic(doc: Document) -> None:
    doc.add_heading("2. 游戏逻辑与产品体验", level=1)
    for text in [
        "开场更长：生成结果会先介绍现实地点、环境氛围、委托触发方式和五站路线预告，而不是只给一句悬念钩子。",
        "身份卡轻量化：人物卡保留探索能力、隐藏目标和提示风格，但不把体验限制为剧本杀。",
        "现实世界链接：每一站都明确“拍什么”，照片会出现 MAIN QUEST 主线标记、HIDDEN 隐藏标记和 NEXT 方向箭头。",
        "推进方式更游戏化：玩家点击照片里的主线任务推进，隐藏任务只解锁额外线索，不强制完成。",
        "路线感更强：剧情页展示 5 个站点、当前站、高亮/完成/待探索状态，以及下一站提示。",
        "完成感更强：线索页在完成全部站点后展示冒险报告和可截图分享卡。",
        "安全边界明确：不要求进入私人区域、打扰路人、拍摄陌生人或接触他人物品。",
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_architecture(doc: Document) -> None:
    doc.add_heading("3. 技术结构", level=1)
    doc.add_paragraph(
        "项目采用 Next.js App Router。客户端状态集中在 AppShell，服务端 API 负责故事生成、身份选择、文字推进、照片分析和照片标记完成。"
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [3200, 6160])
    set_cell_text(table.rows[0].cells[0], "文件 / API", bold=True, fill="F2F4F7")
    set_cell_text(table.rows[0].cells[1], "职责", bold=True, fill="F2F4F7")
    rows = [
        ("src/lib/deepseek-director.ts", "生成城市冒险案卷、身份卡、路线节点，并做结构归一化和安全检查。"),
        ("src/lib/qwen-vision.ts", "调用 Qwen-VL，根据现场照片生成 PhotoOverlay。"),
        ("src/lib/expanded-session.ts", "管理身份选择、文字推进、照片主线/隐藏标记推进。"),
        ("src/app/app-client.tsx", "维护前端状态、底部导航、发起 API 请求，并自动读取最新存档。"),
        ("src/app/api/photo/*", "照片分析与标记完成的后端接口。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0])
        set_cell_text(cells[1], row[1])


def add_ai_flow(doc: Document) -> None:
    doc.add_heading("4. AI 生成与现实世界链接", level=1)
    doc.add_paragraph(
        "DeepSeek 负责叙事导演：先生成案卷、开场、人物和身份卡，再生成五站路线与推理节点。为了避免长 JSON 被截断，当前版本提高了输出 token 上限，并在检测到截断 JSON 时自动用紧凑格式重试。"
    )
    doc.add_paragraph(
        "Qwen-VL 负责照片理解：它只允许基于公共可观察元素生成任务锚点，例如路牌、店招、门口、橱窗、墙面、长椅、地铁口和导视牌。输出会被 normalizePhotoOverlay 再次归一化，兼容 markers、hotspots 和 marks 字段，确保坐标、标记数量和安全文本可控。"
    )
    add_callout(
        doc,
        "关键设计判断",
        "现在的故事不需要强行变成剧本杀。更合适的方向是“城市冒险 + 轻推理 + AR/HUD 拍照任务”，让玩家感觉现实街区正在被游戏化，而不是只是在读案卷。"
    )


def add_page_inventory(doc: Document) -> None:
    doc.add_heading("5. 页面清单", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1800, 2200, 5360])
    for index, text in enumerate(["页面", "路由", "说明"]):
        set_cell_text(table.rows[0].cells[index], text, bold=True, fill="F2F4F7")
    for page in PAGES:
        cells = table.add_row().cells
        set_cell_text(cells[0], page["title"])
        set_cell_text(cells[1], page["route"])
        set_cell_text(cells[2], page["description"])


def add_screenshots(doc: Document, screenshots: dict[str, Path]) -> None:
    doc.add_page_break()
    doc.add_heading("6. 当前页面截图", level=1)
    doc.add_paragraph("以下截图来自当前本地运行版本，覆盖应用底部导航中的六个主要页面。")

    for page in PAGES:
        doc.add_heading(f"{page['title']}（{page['route']}）", level=2)
        doc.add_paragraph(page["description"])
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(screenshots[page["file"]]), width=Inches(3.15))
        caption = doc.add_paragraph(f"图：{page['title']} 当前界面")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_next_steps(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("7. 当前验收点与后续优化", level=1)
    doc.add_paragraph("当前 MVP 已完成的关键验收点：")
    for text in [
        "第一次进入拍照页时，能通过 HUD 示例理解照片游戏化玩法。",
        "照片结果支持主线标记、隐藏标记和方向箭头；隐藏标记不会强制推进主线。",
        "剧情页能看到 5 站路线进度和身份探索加成。",
        "线索页具备冒险完成报告和分享卡结构。",
        "API 失败时使用产品化文案和本地临时线索，不暴露技术错误。",
        "故事生成前端入口包含阶段式 loading，后端有基础质量与安全检查。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("建议后续继续推进：")
    for text in [
        "增加真实照片上传示例和多张照片对比，让 HUD 生成效果更稳定可演示。",
        "把站点路线接入真实地图或步行距离估算，但仍保持公共安全边界。",
        "建立故事吸引力评分：地点密度、开场钩子、隐藏任务趣味性、现实可拍性。",
        "给冒险报告增加可导出图片或海报能力，方便用户分享城市冒险结果。",
        "上线前补充服务端持久化存档，避免 dev 内存存档在重启后丢失。",
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F4F6F9")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    set_cell_margins(cell)
    if fill:
        set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 50, start: int = 100, bottom: int = 50, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


if __name__ == "__main__":
    main()
