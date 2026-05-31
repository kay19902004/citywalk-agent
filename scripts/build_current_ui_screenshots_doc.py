from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/a/Desktop/作品/citywalk-agent")
SCREENSHOT_DIR = ROOT / "docs" / "project-screenshots"
DOCX_OUT = ROOT / "docs" / "CityWalk_当前UI页面截图.docx"

PAGES = [
    ("首页", "/", "01-home.png"),
    ("坐标", "/location", "02-location.png"),
    ("档案", "/dossier", "03-dossier.png"),
    ("任务", "/play", "04-play.png"),
    ("扫描", "/photo", "05-photo.png"),
    ("图鉴", "/clues", "06-clues.png"),
]


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.right_margin = Inches(0.65)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)

    configure_styles(doc)
    add_cover(doc)
    doc.add_page_break()

    for index, (title, route, filename) in enumerate(PAGES):
        if index:
            doc.add_page_break()
        add_page_screenshot(doc, title, route, SCREENSHOT_DIR / filename)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(15, 39, 66)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 18, "0F2742"),
        ("Heading 2", 13, "1976D2"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CityWalk 当前 UI 页面截图")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 39, 66)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("明亮现代 · 城市探索冒险 · 轻游戏 HUD · 收藏图鉴")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(93, 107, 122)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(138, 151, 166)

    doc.add_paragraph()


def add_page_screenshot(doc: Document, title: str, route: str, image_path: Path) -> None:
    heading = doc.add_heading(f"{title}页", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    route_p = doc.add_paragraph()
    route_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    route_run = route_p.add_run(route)
    route_run.font.name = "Arial"
    route_run.font.size = Pt(10)
    route_run.font.color.rgb = RGBColor(93, 107, 122)

    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_p.add_run()
    image_run.add_picture(str(image_path), width=Inches(3.35))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(f"{title}页当前截图")
    caption_run.font.name = "Arial"
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(138, 151, 166)


if __name__ == "__main__":
    main()
