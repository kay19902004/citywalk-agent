from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/a/Desktop/作品/citywalk-agent")
SCREENSHOT_DIR = ROOT / "docs" / "project-screenshots"
DOCX_OUT = ROOT / "docs" / "CityWalk_现实城市冒险Agent_项目说明与页面截图.docx"

PAGES = [
    {
        "title": "首页",
        "route": "/",
        "file": "01-home.png",
        "type": "主入口",
        "description": "开局面板，负责城市选择、冒险模式、环境 Buff、玩家等级和生成任务地图。",
    },
    {
        "title": "任务",
        "route": "/play",
        "file": "02-play.png",
        "type": "主入口",
        "description": "主操作中心，承载当前关卡、路线地图、坐标校准入口、任务档案入口、扫描本站线索和 NPC 提示。",
    },
    {
        "title": "图鉴",
        "route": "/clues",
        "file": "03-clues.png",
        "type": "主入口",
        "description": "线索收集、隐藏彩蛋、收藏图鉴与通关报告入口。",
    },
    {
        "title": "我的",
        "route": "/me",
        "file": "04-me.png",
        "type": "主入口",
        "description": "展示等级、成就、历史冒险和设置入口；当前使用 mock 数据。",
    },
    {
        "title": "坐标校准",
        "route": "/location",
        "file": "05-location.png",
        "type": "子页面",
        "description": "现实校准页，从首页地点卡或任务页进入，用于 GPS、地标、街区类型和现场观察。",
    },
    {
        "title": "任务档案",
        "route": "/dossier",
        "file": "06-dossier.png",
        "type": "子页面",
        "description": "任务档案页，从任务页进入，用于查看故事背景、完成目标、任务奖励和探索职业。",
    },
    {
        "title": "扫描",
        "route": "/photo",
        "file": "07-photo.png",
        "type": "沉浸子页面",
        "description": "照片扫描页，不显示全局底部导航；通过 HUD Demo 或真实照片标记展示 MAIN QUEST、HIDDEN EGG 和 NEXT。",
    },
]


def main() -> None:
    doc = Document()
    configure_document(doc)

    add_cover(doc)
    add_project_overview(doc)
    add_product_flow(doc)
    add_navigation_model(doc)
    add_architecture(doc)
    add_ai_and_safety(doc)
    add_page_table(doc)
    add_screenshots(doc)

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.76)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.76)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(15, 39, 66)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Heading 1", 17, "0F2742", 14, 7),
        ("Heading 2", 13, "1976D2", 10, 5),
        ("Heading 3", 11, "0E8BD8", 7, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "CityWalk 现实城市冒险 Agent"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(138, 151, 166)

    footer = section.footer.paragraphs[0]
    footer.text = "项目说明与页面截图"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(138, 151, 166)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("CityWalk 现实城市冒险 Agent")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 39, 66)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项目说明 · 产品逻辑 · 页面截图")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(25, 118, 210)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  当前 UI：4 主入口 + 子页面任务流")
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(93, 107, 122)

    doc.add_paragraph()
    add_callout(
        doc,
        "一句话概括",
        "CityWalk 是一个把真实街区变成轻量城市冒险游戏的 Agent：用户生成 5 站任务地图，到真实地点拍照，AI 在照片上叠加 MAIN QUEST、HIDDEN EGG 和 NEXT，让城市现场成为可点击的游戏界面。",
    )


def add_project_overview(doc: Document) -> None:
    doc.add_heading("1. 项目总览", level=1)
    doc.add_paragraph(
        "本项目是基于 Next.js App Router 的单人现实城市冒险应用。产品定位不是传统旅游攻略，也不是厚重剧本杀，而是“明亮现代 + 城市探索冒险 + 轻游戏 HUD + 收藏图鉴 + 通关报告”。"
    )
    doc.add_paragraph(
        "用户先在首页选择城市、冒险模式、环境 Buff 和出发时间，生成一条 5 站现实步行任务地图。之后在任务页查看当前关卡与路线进度，进入坐标校准、任务档案或扫描页。到达真实地点后，用户拍照或上传现场照片，AI 会把公共可见元素变成可点击 HUD 标记。"
    )
    doc.add_paragraph(
        "当前 MVP 的核心爽点是照片游戏化：真实街景上浮现 MAIN QUEST 主线标记、HIDDEN EGG 隐藏彩蛋和 NEXT 方向箭头。玩家点击主线标记推进站点，点击隐藏标记收集额外发现，最终在图鉴页查看线索和通关报告。"
    )


def add_product_flow(doc: Document) -> None:
    doc.add_heading("2. 产品与游戏流程", level=1)
    steps = [
        ("开局", "用户在首页选择城市、冒险模式、环境 Buff 和出发时间，点击“生成任务地图”。"),
        ("档案", "生成后可以查看任务档案，阅读背景、目标、奖励，并选择探索职业。"),
        ("任务", "任务页成为主操作中心，展示当前关卡、路线地图、坐标校准入口、任务档案入口和扫描入口。"),
        ("扫描", "用户到真实地点后拍照，系统生成照片 HUD 标记。主线标记推进路线，隐藏彩蛋进入图鉴。"),
        ("图鉴", "所有主线线索、隐藏彩蛋、人物/物件线索在图鉴页汇总，完成后显示通关报告。"),
        ("我的", "我的页沉淀等级、成就、历史冒险和设置入口，后续可接真实用户体系。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1800, 7380])
    set_cell_text(table.rows[0].cells[0], "阶段", bold=True, fill="EAF4FF")
    set_cell_text(table.rows[0].cells[1], "说明", bold=True, fill="EAF4FF")
    for stage, desc in steps:
        cells = table.add_row().cells
        set_cell_text(cells[0], stage)
        set_cell_text(cells[1], desc)


def add_navigation_model(doc: Document) -> None:
    doc.add_heading("3. 当前导航结构", level=1)
    doc.add_paragraph("当前底部导航只保留 4 个主入口，避免移动端拥挤：")
    for text in [
        "首页 /：开局、城市选择、冒险模式、环境 Buff、生成任务地图。",
        "任务 /play：主操作中心，承载路线、当前关卡、坐标校准入口、任务档案入口和扫描入口。",
        "图鉴 /clues：线索收集、隐藏彩蛋、通关报告。",
        "我的 /me：等级、成就、历史冒险、设置。",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_paragraph("以下页面作为子流程，不出现在底部导航：")
    for text in [
        "/location：从首页地点卡或任务页“坐标校准”进入。",
        "/dossier：从任务页“查看任务档案”进入。",
        "/photo：从任务页“扫描本站线索”进入，是沉浸式扫描页，不显示全局底部导航。",
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_architecture(doc: Document) -> None:
    doc.add_heading("4. 技术结构", level=1)
    doc.add_paragraph(
        "前端采用 Next.js App Router。客户端状态集中在 AppShell，页面通过 useCitywalkApp 读取当前城市、地理上下文、生成结果、会话、照片 Overlay 和加载状态。服务端 API 保持独立，前端 UI 调整不改变接口。"
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [3100, 6080])
    set_cell_text(table.rows[0].cells[0], "模块 / 文件", bold=True, fill="EAF4FF")
    set_cell_text(table.rows[0].cells[1], "职责", bold=True, fill="EAF4FF")
    rows = [
        ("src/app/app-client.tsx", "AppShell、前端状态、底部导航、生成/身份/照片/推进 API 调用。"),
        ("src/app/ui.tsx", "PageTitle、GameCard、RouteMapCard、PhotoHudStage、CollectionGrid、ClearReportCard 等 UI 组件。"),
        ("src/lib/deepseek-director.ts", "DeepSeek 故事导演，生成档案、身份卡、5 站路线与节点任务，并做质量归一化。"),
        ("src/lib/qwen-vision.ts", "Qwen-VL 照片识别，生成 HUD 标记、隐藏彩蛋和方向箭头。"),
        ("src/lib/expanded-session.ts", "会话推进、身份选择、主线标记/隐藏标记处理、线索收集。"),
        ("src/app/api/*", "游戏生成、会话读取、身份选择、文字推进、照片分析、照片标记完成等接口。"),
    ]
    for name, desc in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], desc)


def add_ai_and_safety(doc: Document) -> None:
    doc.add_heading("5. AI 能力与安全边界", level=1)
    doc.add_paragraph(
        "DeepSeek 负责生成可玩的城市冒险内容，包括长开场、任务档案、探索职业、5 站路线和每站拍摄目标。生成结果进入前端前会做 normalize / quality guard，确保站点数量、拍摄对象、主线任务和安全边界完整。"
    )
    doc.add_paragraph(
        "Qwen-VL 负责识别用户上传的现场照片，并输出 PhotoOverlay。前端把 Overlay 中的 markers 和 arrows 渲染为可点击 HUD 标记。主线标记完成后推进站点，隐藏彩蛋只加入图鉴，不强制推进。"
    )
    add_callout(
        doc,
        "安全边界",
        "所有线下任务都应保持在公共可见区域：不进入私人区域、不拍摄陌生人、不打扰路人或商户、不接触他人物品。照片识别失败时，前端显示“现场信号有点弱”等产品化文案，并使用临时线索继续体验。",
    )


def add_page_table(doc: Document) -> None:
    doc.add_heading("6. 页面清单", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_widths(table, [1500, 1600, 1600, 4480])
    for idx, text in enumerate(["页面", "路由", "类型", "说明"]):
        set_cell_text(table.rows[0].cells[idx], text, bold=True, fill="EAF4FF")
    for page in PAGES:
        cells = table.add_row().cells
        set_cell_text(cells[0], page["title"])
        set_cell_text(cells[1], page["route"])
        set_cell_text(cells[2], page["type"])
        set_cell_text(cells[3], page["description"])


def add_screenshots(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("7. 当前页面截图", level=1)
    doc.add_paragraph("以下截图来自当前本地运行版本，覆盖 4 个主入口与 3 个子页面。")
    for page in PAGES:
        doc.add_page_break()
        heading = doc.add_heading(f"{page['title']}（{page['route']}）", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph(page["description"])
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(SCREENSHOT_DIR / page["file"]), width=Inches(3.35))
        caption = doc.add_paragraph(f"图：{page['title']}当前界面")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(138, 151, 166)


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9180])
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F0F7FF")
    set_cell_margins(cell, 90, 130, 90, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(25, 118, 210)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    set_cell_margins(cell)
    if fill:
        set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(15, 39, 66)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 60, start: int = 90, bottom: int = 60, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


if __name__ == "__main__":
    main()
