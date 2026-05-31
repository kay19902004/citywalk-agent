from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/a/Desktop/作品/citywalk-agent")
SCREENSHOT_DIR = ROOT / "docs" / "project-screenshots"
OUT = ROOT / "docs" / "CityWalk_现实城市冒险Agent_飞书复制版.docx"


SCREENSHOTS = [
    (
        "首页",
        "作为项目主入口，展示城市冒险氛围、附近探索点、任务入口和“开始冒险”按钮，帮助用户快速进入现实城市任务。",
        "01-home.png",
    ),
    (
        "坐标校准页",
        "支持用户手动填写出发地点、街区类型和现场观察，也可使用 GPS 辅助定位，保证任务能围绕真实地点生成。",
        "05-location.png",
    ),
    (
        "任务档案页",
        "展示故事背景、任务目标、奖励机制和探索职业选择，把 AI 生成的城市剧情包装成可理解、可行动的任务档案。",
        "03-dossier.png",
    ),
    (
        "任务地图页",
        "呈现 5 站任务路线、当前关卡状态和扫描入口，是用户进行城市探索时的主操作中心。",
        "02-play.png",
    ),
    (
        "拍照扫描页",
        "通过 HUD 形式把现场公共可见元素转化为 MAIN QUEST 和 HIDDEN EGG 标记，使真实街景成为可点击的游戏界面。",
        "07-photo.png",
    ),
    (
        "线索图鉴页",
        "汇总主线线索、隐藏彩蛋和图鉴收集进度，增强用户完成任务后的成就感和持续探索动力。",
        "06-clues.png",
    ),
]


def main() -> None:
    missing = [filename for _, _, filename in SCREENSHOTS if not (SCREENSHOT_DIR / filename).exists()]
    if missing:
        raise FileNotFoundError(f"Missing screenshots: {', '.join(missing)}")

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_audience_and_value(doc)
    add_solution(doc)
    add_experience_flow(doc)
    add_screenshots(doc)
    add_innovation(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.76)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.76)

    styles = doc.styles
    normal = styles["Normal"]
    set_run_font(normal.font, 10.5, "0F2742")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.14

    for style_name, size, color, before, after in [
        ("Heading 1", 18, "0B2A4D", 16, 8),
        ("Heading 2", 13, "156CC4", 10, 5),
        ("Heading 3", 11, "123B63", 8, 3),
    ]:
        style = styles[style_name]
        set_run_font(style.font, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("CityWalk 现实城市冒险 Agent")
    format_run(run, 9, "8A97A6")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("项目说明 · 体验流程 · 页面截图")
    format_run(run, 9, "8A97A6")


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("CityWalk 现实城市冒险 Agent")
    format_run(run, 26, "0B2A4D", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项目说明 · 用户体验流程 · 程序页面截图")
    format_run(run, 13, "156CC4", bold=True)

    add_callout(
        doc,
        "文档说明",
        "本文件整合目标人群、项目痛点、用户价值、实现方案、体验流程图、程序页面截图和核心创新点，可直接复制到飞书文档中继续编辑。",
    )


def add_audience_and_value(doc: Document) -> None:
    doc.add_heading("一、目标人群、痛点与用户价值", level=1)
    add_para(doc, "CityWalk 现实城市冒险 Agent 面向喜欢 CityWalk、轻推理、城市打卡、沉浸式互动体验的年轻用户，也适用于文旅景区、商业街区、校园园区等需要低成本互动导览和内容运营的场景。")
    add_para(doc, "当前传统 CityWalk 体验普遍存在三个问题：一是路线和内容同质化，用户往往只能按照固定攻略完成打卡，缺少新鲜感；二是现实地点与叙事内容割裂，城市空间只是背景，难以真正成为故事的一部分；三是用户完成动力不足，路线结束后缺少持续探索、收集和复访机制。")
    add_para(doc, "本项目将真实城市地点转化为可游玩的现实冒险，把用户的当前位置、天气、时间、兴趣偏好和现场照片纳入剧情生成过程。用户不再只是被动浏览路线，而是以探索者身份进入一场由 AI 动态编排的 5 站任务：查看任务档案、选择探索身份、前往真实地点、拍照扫描线索、解锁主线任务和隐藏彩蛋，并在图鉴中沉淀自己的探索成果。")
    add_para(doc, "对个人用户而言，项目提供了更有参与感、游戏感和记忆点的 CityWalk 体验；对文旅、商业街区和校园场景而言，项目可以用较低内容制作成本生成多条可复用、可更新的互动路线，提升空间停留时间、用户参与度和传播分享价值。")


def add_solution(doc: Document) -> None:
    doc.add_heading("二、实现方案", level=1)
    add_para(doc, "项目采用“真实城市地点 + AI 动态剧情 + 拍照识别 + 线索收集”的整体方案，将用户所在街区转化为一场单人现实城市冒险。")
    add_para(doc, "在内容生成层，系统会扫描本地故事素材库，提取剧情母题、场景标签、角色原型、冲突关系、线索接口和结局要素，形成可组合的结构化剧情模块。随后由 DeepSeek 作为剧情导演，根据用户偏好、城市、天气、时间和定位上下文生成任务案卷、角色身份、证物、时间线和 5 个站点任务。")
    add_para(doc, "在现实环境层，系统接入高德地图能力，通过 GPS 和逆地理编码补充城市、地标、周边 POI 与地点类型信息，使任务能够围绕用户真实所处的街区展开。用户也可以手动输入城市和出发地点，保证在定位不可用时仍能生成路线。")
    add_para(doc, "在交互体验层，用户进入任务档案后选择探索职业，并在任务页查看当前站点、路线进度、主线目标和隐藏任务。到达现场后，用户上传或拍摄公共可见元素，系统通过 Qwen-VL 分析照片，将路牌、店招、橱窗、墙面、导视牌等现实锚点转化为可点击的游戏 HUD 标记。点击主线标记可推进剧情，点击隐藏标记可收集额外彩蛋。")
    add_para(doc, "在安全与质量控制层，系统设置了安全兜底和剧情质量门。安全机制会避免引导用户进入私人区域、打扰路人、拍摄陌生人或在恶劣天气、深夜、距离过远等情况下继续前往不适合的地点；剧情质量门会检查站点重复、线索薄弱、矛盾不清、过早泄露真相、拍照任务缺少公共锚点等问题，并在必要时触发局部重写，保证生成内容既可玩又可执行。")


def add_experience_flow(doc: Document) -> None:
    doc.add_heading("三、用户体验流程图", level=1)
    steps = [
        "进入首页",
        "选择地点 / 模式",
        "生成任务地图",
        "查看任务档案",
        "选择探索职业",
        "前往真实站点",
        "拍照扫描",
        "解锁线索 / 彩蛋",
        "图鉴沉淀",
        "通关报告",
    ]
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    set_table_widths(table, [1836, 1836, 1836, 1836, 1836])
    for index, step in enumerate(steps):
        cell = table.rows[index // 5].cells[index % 5]
        set_cell_text(cell, step, bold=True, fill="EAF4FF", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "进入首页 -> 选择地点 / 冒险模式 / 环境 Buff -> 生成任务地图 -> 查看任务档案 -> 选择探索职业 -> 前往真实城市站点 -> 拍照扫描公共可见元素 -> 解锁主线线索 / 隐藏彩蛋 -> 线索图鉴沉淀探索成果 -> 完成路线并生成通关报告",
        color="36536D",
    )


def add_screenshots(doc: Document) -> None:
    doc.add_heading("四、程序页面截图", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [4590, 4590])
    for idx, (title, caption, filename) in enumerate(SCREENSHOTS):
        cell = table.rows[idx // 2].cells[idx % 2]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        shade_cell(cell, "F7FBFF")
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(SCREENSHOT_DIR / filename), width=Inches(2.05))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(f"图{idx + 1}：{title} - {caption}")
        format_run(r, 9, "33536F")


def add_innovation(doc: Document) -> None:
    doc.add_heading("五、核心创新点", level=1)
    points = [
        "项目把静态 CityWalk 升级为可生成、可交互、可收集的现实游戏。传统 CityWalk 主要依赖固定路线和文字介绍，而本项目将路线拆解为任务站点、角色身份、证据线索、隐藏彩蛋和图鉴收集，让城市漫步从“看路线”变成“玩任务”。",
        "项目实现了 AI 剧情与真实环境的动态耦合。系统不是预设一条固定故事线，而是根据用户所在位置、周边环境、天气、时间和偏好实时生成任务，使同一座城市、同一片街区在不同用户和不同时间下都能产生不同的冒险体验。",
        "项目将照片识别引入现实探索玩法。用户拍摄现场公共可见元素后，Qwen-VL 会把照片中的现实锚点转化为主线线索和隐藏任务标记，使真实街景直接参与游戏交互，增强“城市就是游戏场景”的沉浸感。",
        "项目把内容生成与安全、质量约束结合在同一链路中。系统在生成后进行结构校验、剧情质量评估和安全规则过滤，减少路线重复、任务不可执行、线索不清晰或引导风险行为的问题，使 AI 生成内容更适合落地到真实城市空间。",
        "项目具备较强的场景扩展价值。通过更换素材库、地点类型和任务风格，同一套框架可以扩展到文旅导览、商业街区活动、校园迎新、城市文化传播、品牌快闪任务等场景，为线下空间提供可持续更新的互动内容基础设施。",
    ]
    for index, point in enumerate(points, start=1):
        add_para(doc, f"{index}. {point}")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9180])
    cell = table.rows[0].cells[0]
    shade_cell(cell, "EEF7FF")
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    format_run(run, 11, "156CC4", bold=True)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(body)
    format_run(run2, 10.5, "24435F")


def add_para(doc: Document, text: str, color: str = "0F2742") -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    format_run(run, 10.5, color)


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None, align=None) -> None:
    if fill:
        shade_cell(cell, fill)
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    format_run(run, 9.5, "123B63", bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_widths(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            for margin_tag in ["top", "bottom", "left", "right"]:
                tc_mar = tc_pr.find(qn("w:tcMar"))
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                margin = tc_mar.find(qn(f"w:{margin_tag}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_tag}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), "100")
                margin.set(qn("w:type"), "dxa")


def set_run_font(font, size: float, color: str, bold: bool = False) -> None:
    font.name = "Arial"
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    font.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def format_run(run, size: float, color: str, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


if __name__ == "__main__":
    main()
